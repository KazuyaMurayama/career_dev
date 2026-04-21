"""
Update Oza_Career_sheet_20260417.docx -> Oza_Career_sheet_20260421.docx

Changes:
1. Table 0: date -> 2026年4月21日, 稼働開始 -> 2026年5月〜
2. EXECUTIVE SUMMARY (paragraph 4): add Amazon SCM横断 mention
3. Insert new ▎ SCM・業務改革（BPR）・AI適用設計 section (between 成長マネジメント and セキュリティ)
4. Insert new ▶ Amazon SCM achievement bullets (in Amazon 職務経歴 detail section)
5. Table 12: 稼働開始 -> 2026年5月1日〜
"""
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

SRC = '/home/user/career_dev/Oza_Career_sheet_20260417.docx'
DST = '/home/user/career_dev/Oza_Career_sheet_20260421.docx'


def set_run_text(run, text):
    """Replace text content of a w:r element keeping its formatting."""
    # Remove all w:t children
    for t in run.findall(qn('w:t')):
        run.remove(t)
    # Add new w:t
    t = run.makeelement(qn('w:t'), {qn('xml:space'): 'preserve'})
    t.text = text
    run.append(t)


def clone_paragraph(template_p_elem, run_texts):
    """Deepcopy a paragraph element and replace texts of its w:r elements in order.

    run_texts is a list; assigns each text to successive runs.
    If template has more runs than texts, extra runs are removed.
    If template has fewer runs than texts, extras are ignored (error).
    """
    new_p = deepcopy(template_p_elem)
    runs = new_p.findall(qn('w:r'))
    # Remove runs beyond needed
    for r in runs[len(run_texts):]:
        new_p.remove(r)
    # Update kept runs' text
    for r, txt in zip(runs[:len(run_texts)], run_texts):
        # Clear existing w:t and replace
        for t in r.findall(qn('w:t')):
            r.remove(t)
        t_elem = r.makeelement(qn('w:t'), {qn('xml:space'): 'preserve'})
        t_elem.text = txt
        r.append(t_elem)
    return new_p


def main():
    doc = Document(SRC)

    # ========================================================
    # 1. Table 0: update date and 稼働開始
    # ========================================================
    t0_cell = doc.tables[0].rows[0].cells[1]
    # p0: date
    t0_cell.paragraphs[0].runs[0].text = '2026年4月21日'
    # p1: 稼働開始 -> keep run0 '稼働開始', update run1 ': 2026年5月〜'
    t0_cell.paragraphs[1].runs[1].text = ': 2026年5月〜'

    # ========================================================
    # 2. EXECUTIVE SUMMARY paragraph 4
    # ========================================================
    exec_p = doc.paragraphs[4]
    exec_p.runs[0].text = (
        'Amazon Japan シニアデータサイエンティスト（広告事業 新規収益7億円創出・ROAS +30%'
        '／SCM横断プロジェクトで業務工数30%削減・役員会承認）、'
        'エムスリー BIチームリーダー（11名マネジメント・全社売上伸長率No.1）、'
        'STANDARD Senior Manager（生成AIサービスを6ヶ月で年商5,000万円に拡大・大手5社でROI 300%超）。'
        '戦略立案 × 実装 × 組織変革を1人で完結し、コンサルファームの半額以下で同等以上の成果を提供。'
        'SCM・製造・EC/広告・製薬・金融・IT/SaaS・インフラの7領域での導入実績を持つ。'
    )

    # ========================================================
    # 3. Insert ▎ SCM・業務改革（BPR）・AI適用設計 section
    #    Templates: paragraph 13 (▎ section header), paragraph 14 (▶ 3-run bullet)
    #    Insert BEFORE paragraph 13 (▎ セキュリティ・ガバナンス)
    # ========================================================
    section_header_template = doc.paragraphs[13]._element  # ▎ セキュリティ
    bullet3_template = doc.paragraphs[14]._element  # ▶ 3-run bullet (title　 + description)

    new_section_header = clone_paragraph(
        section_header_template,
        ['▎ ', 'SCM・業務改革（BPR）・AI適用設計'],
    )

    new_bullet1 = clone_paragraph(
        bullet3_template,
        [
            '▶ ',
            'SCM領域の判断業務定義・AI/RPA適用マップ策定（Amazon社内横断）　',
            'Amazon日本にて需要予測・物流運用・受注処理を対象に業務棚卸しを実施。'
            '判断知識集約度×ルール化可能性×技術適合度の3軸評価フレームを独自開発し、'
            '判断業務と実行業務を客観的に切り分け。対象部門で業務工数 約30%削減を実現し、'
            '本部長・役員層の承認を獲得。',
        ],
    )

    new_bullet2 = clone_paragraph(
        bullet3_template,
        [
            '▶ ',
            '第三者視点での業務×AI影響レビュー・人員配置ロジック構築　',
            '業務部門／IT／経営企画／物流パートナー（親会社含む）をまたぐステークホルダー間で、'
            '削減根拠と実行性の妥当性を定量検証。役員会提出資料のロジック構築と監修を担当し、'
            '部門再編・人員再配置案の合意形成を主導。',
        ],
    )

    # Insert in order before paragraph 13
    section_header_template.addprevious(new_section_header)
    section_header_template.addprevious(new_bullet1)
    section_header_template.addprevious(new_bullet2)

    # ========================================================
    # 4. Insert Amazon SCM achievement in 職務経歴 detail section
    #    After paragraph 35 ("【施策】金融・動画配信..."), before paragraphs 36-37 blanks
    #    Templates: paragraph 34 (▶ achievement: 4 runs), paragraph 35 (【施策】: 2 runs)
    # ========================================================
    # Find fresh references after previous insertions shifted indices
    # Easier: find target by text
    target_p = None
    for p in doc.paragraphs:
        if p.text.startswith('【施策】金融・動画配信'):
            target_p = p
            break
    assert target_p is not None

    # Template: achievement bullet (paragraph 34-like) with 4 runs
    achievement_template = None
    shisaku_template = None
    for p in doc.paragraphs:
        if p.text.startswith('▶ 新ビジネス領域の立ち上げとグロース'):
            achievement_template = p._element
        if p.text.startswith('【施策】金融・動画配信'):
            shisaku_template = p._element

    new_achievement = clone_paragraph(
        achievement_template,
        [
            '▶ ',
            'SCM／オペレーション業務改革プロジェクト（社内横断）',
            '  【成果】',
            '対象3領域で業務工数 約30%削減、5名相当を戦略業務へ再配置／役員会承認獲得',
        ],
    )

    new_shisaku = clone_paragraph(
        shisaku_template,
        [
            '【施策】',
            '需要予測・物流運用・受注処理の3領域でタスク単位の業務棚卸しを実施。'
            '独自の3軸評価フレーム（判断知識集約度×ルール化可能性×AI/RPA技術適合度）で'
            '判断業務と実行業務を客観的に切り分け。現場／部長／IT／経営企画／親会社物流パートナーを跨ぐ'
            '第三者視点でステークホルダー合意形成を主導。役員会提出資料の削減根拠と実行性妥当性を定量検証・監修。',
        ],
    )

    # Insert AFTER shisaku_template (after paragraph 35)
    shisaku_template.addnext(new_shisaku)  # inserted: [35, new_shisaku, 36...]
    shisaku_template.addnext(new_achievement)  # inserted: [35, new_achievement, new_shisaku, 36...]

    # ========================================================
    # 5. Table 12: 稼働条件 - update 稼働開始 to 2026年5月1日〜
    # ========================================================
    t12_cond_cell = doc.tables[12].rows[0].cells[1]
    # p1: '稼働開始　2026年6月1日〜'
    p1 = t12_cond_cell.paragraphs[1]
    # Rebuild text while preserving first run's format
    full_text = p1.text
    new_text = full_text.replace('2026年6月1日〜', '2026年5月1日〜')
    # Replace all run texts: put new_text into first run, clear others
    if p1.runs:
        p1.runs[0].text = new_text
        for r in p1.runs[1:]:
            r.text = ''

    # Save
    doc.save(DST)
    print(f"Saved: {DST}")


if __name__ == '__main__':
    main()
